from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submission" / "DriveLens_复赛参赛方案.md"
OUTPUT = ROOT / "submission" / "DriveLens_复赛参赛方案.docx"
HERO = ROOT / "public" / "og-v2.png"

BLUE = "1677FF"
CYAN = "13C2C2"
NAVY = "0B2545"
INK = "172033"
MUTED = "5F6B7A"
LIGHT_BLUE = "EAF2FF"
LIGHT_CYAN = "E8F7F6"
LIGHT_GRAY = "F4F6F9"
BORDER = "D8E0EB"
WHITE = "FFFFFF"
RED = "B42318"


def set_font(run, size: float | None = None, bold: bool | None = None,
             color: str | None = None, italic: bool | None = None,
             ascii_font: str = "Calibri", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120,
                     bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must sum to 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_border(paragraph, color: str, size: int = 10, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run("DriveLens  |  飞书 AI 先锋未来人才大赛复赛方案")
        set_font(hr, 8.5, color=MUTED)
        set_paragraph_border(hp, BORDER, size=5, space=4)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("DriveLens · 2026-08-16  |  ")
        set_font(fr, 8.5, color=MUTED)
        add_field(fp, " PAGE ")


def add_inline_runs(paragraph, text: str, base_size: float = 11) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, base_size, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, base_size, bold=True, color=NAVY)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, base_size - 0.5, color=BLUE, ascii_font="Consolas", east_asia="Microsoft YaHei")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, base_size, color=INK)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("飞书 AI 先锋未来人才大赛 · 复赛参赛方案")
    set_font(r, 11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DriveLens")
    set_font(r, 30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("无人车异常行为诊断工具箱")
    set_font(r, 16, bold=True, color=CYAN)

    if HERO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        inline_shape = p.add_run().add_picture(str(HERO), width=Inches(6.5))
        # Keep the cover visual accessible when the DOCX is inspected by screen readers.
        inline_shape._inline.docPr.set("title", "DriveLens 项目封面")
        inline_shape._inline.docPr.set(
            "descr",
            "DriveLens 无人车异常行为诊断工具箱的项目封面图",
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("可回放 · 可反驳 · 可协同")
    set_font(r, 14, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("企业命题：佑驾创新无人车异常行为诊断工具箱")
    set_font(r, 10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("版本 v1.0.0  |  2026 年 8 月 16 日")
    set_font(r, 10, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def column_widths(rows: list[list[str]]) -> list[int]:
    count = len(rows[0])
    if count == 2:
        left_max = max(len(row[0]) for row in rows)
        return [2200, 7160] if left_max <= 14 else [3000, 6360]
    if count == 3:
        return [2200, 3380, 3780]
    if count == 4:
        return [1900, 1700, 2600, 3160]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ridx, row in enumerate(rows):
        for cidx, text in enumerate(row):
            cell = table.cell(ridx, cidx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ridx == 0 or cidx == 0 and len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, text, 9.5 if len(rows[0]) >= 4 else 10)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
            elif ridx % 2 == 0:
                set_cell_shading(cell, "FAFBFD")
    set_repeat_table_header(table.rows[0])
    set_table_borders(table)
    set_table_geometry(table, column_widths(rows))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0E1726")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, 9, color=WHITE, ascii_font="Consolas", east_asia="Microsoft YaHei")
    set_table_borders(table, color="243247", size=5)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next((idx for idx, line in enumerate(lines) if line.startswith("## 参赛方案信息卡")), 0)
    lines = lines[start:]
    paragraph_buffer: list[str] = []
    table_rows: list[list[str]] = []
    code_lines: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = "".join(item.strip() for item in paragraph_buffer)
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            add_inline_runs(p, text)
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            clean = [row for idx, row in enumerate(table_rows) if idx != 1 or not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in row)]
            add_markdown_table(doc, clean)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_paragraph()
            flush_table()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            continue
        flush_table()
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("#### "):
            flush_paragraph()
            doc.add_heading(line[5:].strip(), level=3)
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=1)
            continue
        match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, match.group(2))
            continue
        if line.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            continue
        paragraph_buffer.append(line)

    flush_paragraph()
    flush_table()
    if code_lines:
        add_code_block(doc, code_lines)


def add_final_callout(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DriveLens 交付的不是一段 AI 答案，")
    set_font(r, 13, bold=True, color=NAVY)
    r = p.add_run("而是一条能被工程师验证、反驳和沉淀复用的证据链。")
    set_font(r, 13, bold=True, color=BLUE)


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    add_final_callout(doc)

    doc.core_properties.title = "DriveLens：无人车异常行为诊断工具箱"
    doc.core_properties.subject = "飞书 AI 先锋未来人才大赛复赛参赛方案"
    doc.core_properties.author = "殷该很蔡队 · 殷昊颉、蔡昊哲、陈硕涵"
    doc.core_properties.keywords = "DriveLens, 无人车, 异常诊断, 飞书AI, 证据链"
    doc.core_properties.comments = "Generated from the versioned submission source."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
